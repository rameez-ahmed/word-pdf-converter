#!/usr/bin/env python3
"""
convert.py — Word↔PDF conversion script
Called by Node.js via execFile:
  python3 convert.py word-to-pdf <pageSize> <orientation> <marginMm> <inFile> <outFile>
  python3 convert.py pdf-to-word <inFile> <outFile>
"""

import sys
import os

# Add user pip packages to path
sys.path.insert(0, os.path.expanduser('~/.local/lib/python3.9/site-packages'))

def word_to_pdf(in_file, out_file, page_size='A4', orientation='portrait', margin_mm=20):
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from reportlab.lib.pagesizes import A4, LETTER, LEGAL, landscape, portrait as rl_portrait
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, Image, HRFlowable,
                                     PageBreak, ListFlowable, ListItem)
    from reportlab.platypus.flowables import KeepTogether
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    import io
    import base64
    import re

    # Page size
    size_map = {
        'A4':     A4,
        'LETTER': LETTER,
        'LEGAL':  LEGAL,
    }
    page_size_upper = page_size.upper()
    base_size = size_map.get(page_size_upper, A4)
    if orientation.lower() == 'landscape':
        page_size_rl = landscape(base_size)
    else:
        page_size_rl = rl_portrait(base_size)

    margin = float(margin_mm) * mm

    doc = Document(in_file)

    # Build PDF
    pdf_buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        pdf_buf,
        pagesize=page_size_rl,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    def make_style(name, parent='Normal', **kwargs):
        if name in styles:
            return styles[name]
        s = ParagraphStyle(name, parent=styles[parent], **kwargs)
        styles.add(s)
        return s

    style_normal = make_style('DocNormal', fontSize=11, leading=16,
                               spaceAfter=6, fontName='Helvetica')
    style_h1 = make_style('DocH1', fontSize=22, leading=28, fontName='Helvetica-Bold',
                           spaceBefore=12, spaceAfter=8, textColor=colors.HexColor('#1a1a1a'))
    style_h2 = make_style('DocH2', fontSize=17, leading=22, fontName='Helvetica-Bold',
                           spaceBefore=10, spaceAfter=6, textColor=colors.HexColor('#1a1a1a'))
    style_h3 = make_style('DocH3', fontSize=14, leading=18, fontName='Helvetica-Bold',
                           spaceBefore=8, spaceAfter=4)
    style_h4 = make_style('DocH4', fontSize=12, leading=16, fontName='Helvetica-Bold',
                           spaceBefore=6, spaceAfter=4)
    style_h5 = make_style('DocH5', fontSize=11, leading=15, fontName='Helvetica-BoldOblique',
                           spaceBefore=4, spaceAfter=2)
    style_h6 = make_style('DocH6', fontSize=10, leading=14, fontName='Helvetica-Oblique',
                           spaceBefore=4, spaceAfter=2)

    def get_align(para):
        try:
            a = para.alignment
            if a == WD_ALIGN_PARAGRAPH.CENTER:  return TA_CENTER
            if a == WD_ALIGN_PARAGRAPH.RIGHT:   return TA_RIGHT
            if a == WD_ALIGN_PARAGRAPH.JUSTIFY: return TA_JUSTIFY
        except:
            pass
        return TA_LEFT

    def runs_to_markup(para):
        """Convert paragraph runs to ReportLab markup string."""
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            # Escape XML special chars
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Apply formatting
            if run.bold:       text = f'<b>{text}</b>'
            if run.italic:     text = f'<i>{text}</i>'
            if run.underline:  text = f'<u>{text}</u>'

            # Font color
            try:
                if run.font.color and run.font.color.rgb:
                    hex_color = str(run.font.color.rgb)
                    text = f'<font color="#{hex_color}">{text}</font>'
            except:
                pass

            # Font size
            try:
                if run.font.size and run.font.size.pt:
                    sz = int(run.font.size.pt)
                    text = f'<font size="{sz}">{text}</font>'
            except:
                pass

            parts.append(text)
        return ''.join(parts)

    def para_to_flowable(para):
        style_name = para.style.name if para.style else 'Normal'
        markup = runs_to_markup(para)
        align  = get_align(para)

        if not markup.strip():
            return Spacer(1, 6)

        # Pick style
        if   'Heading 1' in style_name: base = style_h1
        elif 'Heading 2' in style_name: base = style_h2
        elif 'Heading 3' in style_name: base = style_h3
        elif 'Heading 4' in style_name: base = style_h4
        elif 'Heading 5' in style_name: base = style_h5
        elif 'Heading 6' in style_name: base = style_h6
        else:                            base = style_normal

        ps = ParagraphStyle(
            'inline_' + str(id(para)),
            parent=base,
            alignment=align
        )
        try:
            return Paragraph(markup, ps)
        except Exception:
            # Fallback: plain text
            plain = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return Paragraph(plain, ps)

    def table_to_flowable(tbl):
        """Convert a python-docx Table to a ReportLab Table."""
        data = []
        for row in tbl.rows:
            row_data = []
            for cell in row.cells:
                # Get all paragraphs in cell
                cell_parts = []
                for p in cell.paragraphs:
                    markup = runs_to_markup(p)
                    if markup.strip():
                        try:
                            cell_parts.append(Paragraph(markup, style_normal))
                        except:
                            plain = p.text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                            cell_parts.append(Paragraph(plain, style_normal))
                    else:
                        cell_parts.append(Spacer(1, 3))
                row_data.append(cell_parts if cell_parts else [Paragraph('', style_normal)])
            data.append(row_data)

        if not data:
            return None

        # Available width
        available_w = page_size_rl[0] - 2 * margin
        col_count = max(len(r) for r in data)
        col_w = available_w / col_count if col_count else available_w

        tbl_style = TableStyle([
            ('BACKGROUND',  (0, 0), (-1, 0),  colors.HexColor('#f0f0f0')),
            ('FONTNAME',    (0, 0), (-1, 0),  'Helvetica-Bold'),
            ('FONTSIZE',    (0, 0), (-1, -1), 10),
            ('GRID',        (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
            ('VALIGN',      (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING',(0, 0), (-1, -1), 6),
            ('TOPPADDING',  (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING',(0,0), (-1, -1), 4),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
        ])

        col_widths = [col_w] * col_count
        rl_table = Table(data, colWidths=col_widths, repeatRows=1)
        rl_table.setStyle(tbl_style)
        return rl_table

    # ── Build flowables from document ──
    story = []
    elements = list(doc.element.body)

    # We iterate doc.paragraphs and doc.tables in order using XML element iteration
    from docx.oxml.ns import qn

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # It's a paragraph
            from docx.text.paragraph import Paragraph as DocxParagraph
            para = DocxParagraph(child, doc)
            story.append(para_to_flowable(para))

        elif tag == 'tbl':
            # It's a table
            from docx.table import Table as DocxTable
            tbl = DocxTable(child, doc)
            ft = table_to_flowable(tbl)
            if ft:
                story.append(ft)
            story.append(Spacer(1, 8))

        elif tag == 'sectPr':
            # Section properties — ignore
            pass

    if not story:
        story = [Paragraph('(Empty document)', style_normal)]

    pdf_doc.build(story)
    pdf_bytes = pdf_buf.getvalue()

    with open(out_file, 'wb') as f:
        f.write(pdf_bytes)

    print(f"word-to-pdf: wrote {len(pdf_bytes)} bytes to {out_file}")


def pdf_to_word(in_file, out_file):
    """
    Extract text from PDF page by page and build a .docx file.
    Uses pdfminer.six for text extraction (best pure-Python option).
    Falls back to basic extraction if not available.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # Set margins
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = Mm(20)
    section.top_margin  = section.bottom_margin = Mm(20)

    # Style the default paragraph font
    style = doc.styles['Normal']
    font  = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    text_extracted = False

    # Try pdfminer first (best quality)
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer, LTTextLine, LTChar, LTAnon, LAParams

        laparams = LAParams(line_margin=0.5, word_margin=0.1)
        page_num = 0

        for page_layout in extract_pages(in_file, laparams=laparams):
            page_num += 1
            if page_num > 1:
                doc.add_page_break()

            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    text = element.get_text().strip()
                    if text:
                        para = doc.add_paragraph(text)
                        para.style = doc.styles['Normal']

        text_extracted = True

    except ImportError:
        pass

    # Try pypdf as fallback
    if not text_extracted:
        try:
            import pypdf
            reader = pypdf.PdfReader(in_file)
            for i, page in enumerate(reader.pages):
                if i > 0:
                    doc.add_page_break()
                text = page.extract_text() or ''
                lines = [l.strip() for l in text.split('\n') if l.strip()]
                for line in lines:
                    doc.add_paragraph(line)
            text_extracted = True
        except ImportError:
            pass

    # Last fallback — basic PyPDF2
    if not text_extracted:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(in_file)
            for i, page in enumerate(reader.pages):
                if i > 0:
                    doc.add_page_break()
                text = page.extract_text() or ''
                lines = [l.strip() for l in text.split('\n') if l.strip()]
                for line in lines:
                    doc.add_paragraph(line)
            text_extracted = True
        except ImportError:
            pass

    if not text_extracted:
        doc.add_paragraph('Could not extract text from this PDF. '
                          'This may be a scanned/image-based PDF.')

    doc.save(out_file)
    size = os.path.getsize(out_file)
    print(f"pdf-to-word: wrote {size} bytes to {out_file}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: convert.py <mode> [options] <inFile> <outFile>")
        sys.exit(1)

    mode    = sys.argv[1]
    in_file = sys.argv[-2]
    out_file= sys.argv[-1]

    if mode == 'word-to-pdf':
        page_size   = sys.argv[2] if len(sys.argv) > 4 else 'A4'
        orientation = sys.argv[3] if len(sys.argv) > 5 else 'portrait'
        margin_mm   = sys.argv[4] if len(sys.argv) > 6 else '20'
        word_to_pdf(in_file, out_file, page_size, orientation, margin_mm)

    elif mode == 'pdf-to-word':
        pdf_to_word(in_file, out_file)

    else:
        print(f"Unknown mode: {mode}")
        sys.exit(1)
